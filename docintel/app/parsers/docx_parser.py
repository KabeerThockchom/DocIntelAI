import os
import time
from typing import Dict, List, Any, Optional, BinaryIO, Tuple
import docx
from io import BytesIO

from app.parsers.base_parser import BaseDocumentParser
from app.parsers.ocr import OCRProcessor
from app.chunking.models import ProcessedDocument
from app.chunking.chunker import DocumentChunker
from app.utils.logging import log_step, Timer


class DocxParser(BaseDocumentParser):
    """Parser for DOCX documents."""
    
    def __init__(self, chunker: Optional[DocumentChunker] = None):
        """
        Initialize DOCX parser.
        
        Args:
            chunker: Document chunker instance (optional)
        """
        super().__init__()
        self.chunker = chunker or DocumentChunker()
        self.ocr_processor = OCRProcessor()
    
    def parse(
        self, 
        file_path: str, 
        filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """
        Parse a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            filename: Original filename (uses basename of file_path if not provided)
            metadata: Additional metadata
            
        Returns:
            ProcessedDocument object with extracted content and metadata
        """
        with Timer("DOCX Parsing"):
            # Get filename if not provided
            if not filename:
                filename = os.path.basename(file_path)
                
            log_step("DOCX Parsing", f"Parsing DOCX file: {filename}")
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            # Prepare metadata
            doc_metadata = self.prepare_metadata(filename, file_size, metadata)
            
            # Check if force_ocr is set in metadata
            force_ocr = metadata.get("force_ocr", False) if metadata else False
            
            # Start timer for processing
            start_time = time.time()
            
            try:
                use_ocr = False
                text_by_section = {}
                
                # If not forcing OCR, try simple DOCX parsing first
                if not force_ocr:
                    text_by_section, is_complex = self._extract_text_from_docx(file_path)
                    use_ocr = is_complex
                else:
                    log_step("DOCX Parsing", "Force OCR is enabled, skipping normal text extraction")
                    is_complex = True
                    use_ocr = True
                
                # Use OCR if needed (either forced or complex document)
                if use_ocr:
                    log_step("DOCX Parsing", "Using OCR for DOCX processing")
                    ocr_results = self.ocr_processor.process_file(file_path, "docx")
                    
                    # Convert OCR results to text by section format
                    text_by_section = {}
                    if ocr_results:
                        text_by_section = {
                            f"section_{result['page_number']}": {
                                "text": result["text"],
                                "is_ocr": True,
                                "heading": f"Section {result['page_number']}"
                            }
                            for result in ocr_results
                            if result and result.get("text")
                        }
                    
                    # If OCR failed to extract any text, try one more time with higher quality
                    if not text_by_section:
                        log_step("DOCX Parsing", "OCR failed, retrying with higher quality settings", level="warning")
                        self.ocr_processor = OCRProcessor(max_workers=2)  # Reduce workers but increase quality
                        ocr_results = self.ocr_processor.process_file(file_path, "docx")
                        if ocr_results:
                            text_by_section = {
                                f"section_{result['page_number']}": {
                                    "text": result["text"],
                                    "is_ocr": True,
                                    "heading": f"Section {result['page_number']}"
                                }
                                for result in ocr_results
                                if result and result.get("text")
                            }
                
                # No text extracted
                if not text_by_section:
                    log_step("DOCX Parsing", "No text extracted from DOCX", level="warning")
                    return ProcessedDocument(
                        document_id=self.document_id,
                        filename=filename,
                        file_type="docx",
                        file_size=file_size,
                        total_pages=0,
                        total_chunks=0,
                        chunks=[],
                        processing_time=time.time() - start_time,
                        is_complex=is_complex
                    )
                
                # Process each section
                all_chunks = []
                
                for section_id, section_data in text_by_section.items():
                    section_text = section_data["text"]
                    is_ocr = section_data.get("is_ocr", False)
                    section_heading = section_data.get("heading", "")
                    
                    # Skip empty sections
                    if not section_text:
                        continue
                    
                    # Add section info to metadata
                    section_metadata = doc_metadata.copy()
                    section_metadata["section_id"] = section_id
                    section_metadata["is_ocr"] = is_ocr
                    if section_heading:
                        section_metadata["section_heading"] = section_heading
                    
                    # Chunk the section
                    section_chunks = self.chunker.chunk_document(
                        section_text,
                        section_metadata,
                        use_headings=True,
                        is_ocr=is_ocr
                    )
                    
                    all_chunks.extend(section_chunks)
                
                # Estimate number of pages
                total_pages = len(text_by_section)
                
                # Create processed document
                processed_doc = ProcessedDocument(
                    document_id=self.document_id,
                    filename=filename,
                    file_type="docx",
                    file_size=file_size,
                    total_pages=total_pages,
                    total_chunks=len(all_chunks),
                    chunks=all_chunks,
                    processing_time=time.time() - start_time,
                    is_complex=is_complex
                )
                
                log_step("DOCX Parsing", f"Completed parsing DOCX with {total_pages} sections and {len(all_chunks)} chunks")
                return processed_doc
                
            except Exception as e:
                log_step("DOCX Parsing", f"Error parsing DOCX: {str(e)}", level="error")
                raise
    
    def parse_stream(
        self, 
        file_stream: BinaryIO,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """
        Parse a DOCX file from a byte stream.
        
        Args:
            file_stream: File-like object containing DOCX data
            filename: Original filename
            metadata: Additional metadata
            
        Returns:
            ProcessedDocument object with extracted content and metadata
        """
        with Timer("DOCX Stream Parsing"):
            log_step("DOCX Parsing", f"Parsing DOCX stream: {filename}")
            
            # Get file size
            file_stream.seek(0, 2)  # Seek to end
            file_size = file_stream.tell()  # Get position (size)
            file_stream.seek(0)  # Reset to beginning
            
            # Prepare metadata
            doc_metadata = self.prepare_metadata(filename, file_size, metadata)
            
            # Check if force_ocr is set in metadata
            force_ocr = metadata.get("force_ocr", False) if metadata else False
            
            # Start timer for processing
            start_time = time.time()
            
            try:
                # Create in-memory file
                docx_data = file_stream.read()
                memory_stream = BytesIO(docx_data)
                
                use_ocr = False
                text_by_section = {}
                
                # If not forcing OCR, try simple DOCX parsing first
                if not force_ocr:
                    text_by_section, is_complex = self._extract_text_from_docx_stream(memory_stream)
                    use_ocr = is_complex
                else:
                    log_step("DOCX Parsing", "Force OCR is enabled for stream, skipping normal text extraction")
                    is_complex = True
                    use_ocr = True
                
                # Use OCR if needed (either forced or complex document)
                if use_ocr:
                    log_step("DOCX Parsing", "Using OCR for DOCX stream processing")
                    # Reset file stream
                    file_stream.seek(0)
                    ocr_results = self.ocr_processor.process_stream(file_stream, "docx", filename)
                    
                    # Convert OCR results to text by section format
                    text_by_section = {}
                    if ocr_results:
                        text_by_section = {
                            f"section_{result['page_number']}": {
                                "text": result["text"],
                                "is_ocr": True,
                                "heading": f"Section {result['page_number']}"
                            }
                            for result in ocr_results
                            if result and result.get("text")
                        }
                    
                    # If OCR failed to extract any text, try one more time with higher quality
                    if not text_by_section:
                        log_step("DOCX Parsing", "OCR failed, retrying with higher quality settings", level="warning")
                        self.ocr_processor = OCRProcessor(max_workers=2)  # Reduce workers but increase quality
                        ocr_results = self.ocr_processor.process_stream(file_stream, "docx", filename)
                        if ocr_results:
                            text_by_section = {
                                f"section_{result['page_number']}": {
                                    "text": result["text"],
                                    "is_ocr": True,
                                    "heading": f"Section {result['page_number']}"
                                }
                                for result in ocr_results
                                if result and result.get("text")
                            }
                
                # No text extracted
                if not text_by_section:
                    log_step("DOCX Parsing", "No text extracted from DOCX", level="warning")
                    return ProcessedDocument(
                        document_id=self.document_id,
                        filename=filename,
                        file_type="docx",
                        file_size=file_size,
                        total_pages=0,
                        total_chunks=0,
                        chunks=[],
                        processing_time=time.time() - start_time,
                        is_complex=is_complex
                    )
                
                # Process each section
                all_chunks = []
                
                for section_id, section_data in text_by_section.items():
                    section_text = section_data["text"]
                    is_ocr = section_data.get("is_ocr", False)
                    section_heading = section_data.get("heading", "")
                    
                    # Skip empty sections
                    if not section_text:
                        continue
                    
                    # Add section info to metadata
                    section_metadata = doc_metadata.copy()
                    section_metadata["section_id"] = section_id
                    section_metadata["is_ocr"] = is_ocr
                    if section_heading:
                        section_metadata["section_heading"] = section_heading
                    
                    # Chunk the section
                    section_chunks = self.chunker.chunk_document(
                        section_text,
                        section_metadata,
                        use_headings=True,
                        is_ocr=is_ocr
                    )
                    
                    all_chunks.extend(section_chunks)
                
                # Estimate number of pages
                total_pages = len(text_by_section)
                
                # Create processed document
                processed_doc = ProcessedDocument(
                    document_id=self.document_id,
                    filename=filename,
                    file_type="docx",
                    file_size=file_size,
                    total_pages=total_pages,
                    total_chunks=len(all_chunks),
                    chunks=all_chunks,
                    processing_time=time.time() - start_time,
                    is_complex=is_complex
                )
                
                log_step("DOCX Parsing", f"Completed parsing DOCX with {total_pages} sections and {len(all_chunks)} chunks")
                return processed_doc
                
            except Exception as e:
                log_step("DOCX Parsing", f"Error parsing DOCX stream: {str(e)}", level="error")
                raise
    
    def _extract_text_from_docx(self, file_path: str) -> Tuple[Dict[str, Dict[str, Any]], bool]:
        """
        Extract text from DOCX using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (text by section dictionary, is complex document flag)
        """
        text_by_section = {}
        current_heading = "Document"
        current_section_id = "section_1"
        section_text = []
        is_complex = False
        
        try:
            doc = docx.Document(file_path)
            
            # Process each paragraph
            for i, para in enumerate(doc.paragraphs):
                # Extract paragraph style and text
                style_name = para.style.name if para.style else ""
                text = para.text.strip()
                
                # Check if paragraph is a heading
                if "Heading" in style_name or style_name in ["Title", "Subtitle"]:
                    # If we already have content in the current section, save it
                    if section_text:
                        text_by_section[current_section_id] = {
                            "text": "\n".join(section_text),
                            "heading": current_heading,
                            "is_ocr": False
                        }
                        
                        # Reset section text
                        section_text = []
                    
                    # Update current heading and section ID
                    current_heading = text
                    current_section_id = f"section_{len(text_by_section) + 1}"
                
                # Add paragraph text to current section
                if text:
                    section_text.append(text)
            
            # Add the last section if it has content
            if section_text:
                text_by_section[current_section_id] = {
                    "text": "\n".join(section_text),
                    "heading": current_heading,
                    "is_ocr": False
                }
            
            # Check if document is complex (needs OCR)
            total_text_length = sum(len(section["text"]) for section in text_by_section.values())
            if total_text_length < 100 or len(text_by_section) == 0:
                is_complex = True
            
            return text_by_section, is_complex
            
        except Exception as e:
            log_step("DOCX Parsing", f"Error in simple DOCX parsing: {str(e)}", level="warning")
            return {}, True
    
    def _extract_text_from_docx_stream(self, memory_stream: BytesIO) -> Tuple[Dict[str, Dict[str, Any]], bool]:
        """
        Extract text from DOCX stream using python-docx.
        
        Args:
            memory_stream: BytesIO object containing the DOCX
            
        Returns:
            Tuple of (text by section dictionary, is complex document flag)
        """
        text_by_section = {}
        current_heading = "Document"
        current_section_id = "section_1"
        section_text = []
        is_complex = False
        
        try:
            # Reset stream position
            memory_stream.seek(0)
            
            doc = docx.Document(memory_stream)
            
            # Process each paragraph
            for i, para in enumerate(doc.paragraphs):
                # Extract paragraph style and text
                style_name = para.style.name if para.style else ""
                text = para.text.strip()
                
                # Check if paragraph is a heading
                if "Heading" in style_name or style_name in ["Title", "Subtitle"]:
                    # If we already have content in the current section, save it
                    if section_text:
                        text_by_section[current_section_id] = {
                            "text": "\n".join(section_text),
                            "heading": current_heading,
                            "is_ocr": False
                        }
                        
                        # Reset section text
                        section_text = []
                    
                    # Update current heading and section ID
                    current_heading = text
                    current_section_id = f"section_{len(text_by_section) + 1}"
                
                # Add paragraph text to current section
                if text:
                    section_text.append(text)
            
            # Add the last section if it has content
            if section_text:
                text_by_section[current_section_id] = {
                    "text": "\n".join(section_text),
                    "heading": current_heading,
                    "is_ocr": False
                }
            
            # Check if document is complex (needs OCR)
            total_text_length = sum(len(section["text"]) for section in text_by_section.values())
            if total_text_length < 100 or len(text_by_section) == 0:
                is_complex = True
            
            return text_by_section, is_complex
            
        except Exception as e:
            log_step("DOCX Parsing", f"Error in simple DOCX parsing from stream: {str(e)}", level="warning")
            return {}, True